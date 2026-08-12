"""Unit tests for the DOCX usage-guide renderer.

Run with:
    python -m unittest discover summarize-product-usage/tests
"""

from __future__ import annotations

import base64
import importlib.util
import json
import re
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


SCRIPT = Path(__file__).resolve().parents[1] / "scripts" / "render_usage_summary_docx.py"
FIXTURES = Path(__file__).resolve().parent / "fixtures"

# Minimal valid 1x1 PNG, generated at test time so no binary is committed.
MINIMAL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQ"
    "AAAABJRU5ErkJggg=="
)


def load_renderer():
    spec = importlib.util.spec_from_file_location("render_usage_summary_docx", SCRIPT)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def heading_texts(doc: Document) -> list[str]:
    return [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]


def all_texts(doc: Document) -> list[str]:
    """Body paragraphs plus table-cell paragraphs (callouts live in tables)."""
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
    return texts


class RendererTestBase(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.renderer = load_renderer()
        cls.tmp = tempfile.TemporaryDirectory()
        cls.tmpdir = Path(cls.tmp.name)
        cls.png = cls.tmpdir / "sample.png"
        cls.png.write_bytes(MINIMAL_PNG)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.tmp.cleanup()

    def load_fixture(self, name: str) -> dict:
        return json.loads((FIXTURES / name).read_text(encoding="utf-8"))

    @staticmethod
    def replace_image_paths(value, old: str, new: str):
        """Recursively replace fixture image paths with the test-time PNG."""
        if isinstance(value, dict):
            return {k: RendererTestBase.replace_image_paths(v, old, new) for k, v in value.items()}
        if isinstance(value, list):
            return [RendererTestBase.replace_image_paths(v, old, new) for v in value]
        if value == old:
            return new
        return value

    def render_data(self, data: dict, name: str = "out.docx"):
        output = self.tmpdir / name
        self.renderer.render(data, output, FIXTURES)
        return output

    def render_fixture(self, name: str, patch_png: bool = True, output_name: str | None = None):
        data = self.load_fixture(name)
        if patch_png:
            data = self.replace_image_paths(data, "images/sample.png", str(self.png))
            data = self.replace_image_paths(data, "images/legacy.png", str(self.png))
        return self.render_data(data, output_name or f"{name}.docx")

    def open_docx(self, path: Path) -> Document:
        return Document(str(path))


class TestLegacy(RendererTestBase):
    def test_legacy_generates_and_reopens(self) -> None:
        out = self.render_fixture("legacy.json")
        doc = self.open_docx(out)
        paras = [p.text for p in doc.paragraphs]
        self.assertIn("内容趋势分析平台 公司内部同事使用说明", paras[0])
        self.assertIn("模块使用说明", paras)
        self.assertIn("1. 进入左侧菜单「趋势洞察」。", paras)
        self.assertIn("常见问题", paras)
        self.assertIn("最常用的几件事", paras)
        self.assertNotIn("问题反馈模板", paras)

    def test_legacy_module_sections_and_screenshots(self) -> None:
        out = self.render_fixture("legacy.json")
        doc = self.open_docx(out)
        paras = [p.text for p in doc.paragraphs]
        self.assertIn("什么时候用", paras)
        self.assertIn("重点看什么", paras)
        self.assertGreaterEqual(len(doc.inline_shapes), 1)
        self.assertTrue(any("内容生成页面" in p for p in paras))

    def test_legacy_figure_numbering_global_and_continuous(self) -> None:
        data = self.load_fixture("legacy.json")
        png2 = self.tmpdir / "sample2.png"
        png2.write_bytes(MINIMAL_PNG)
        data["modules"][0]["screenshots"] = [{"path": str(png2), "caption": "趋势页面"}]
        data["modules"][1]["screenshots"] = [{"path": str(self.png), "caption": "内容生成页面"}]
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertTrue(any("图 1：趋势页面" in p for p in paras))
        self.assertTrue(any("图 2：内容生成页面" in p for p in paras))

    def test_legacy_module_fallback_fields(self) -> None:
        out = self.render_fixture("legacy.json")
        doc = self.open_docx(out)
        paras = [p.text for p in doc.paragraphs]
        for heading in ("适用场景", "操作步骤", "重点结果", "注意事项"):
            self.assertIn(heading, paras)

    def test_legacy_module_object_step(self) -> None:
        data = self.load_fixture("legacy.json")
        data["modules"][1]["steps"] = [
            "选择趋势方向。",
            {
                "action": "点击「开始生成」。",
                "expected_result": "任务状态变为「处理中」。",
                "warning": "处理中不要重复提交。",
            },
        ]
        texts = all_texts(self.open_docx(self.render_data(data)))
        joined = "\n".join(texts)
        self.assertIn("2. 点击「开始生成」。", joined)
        self.assertIn("完成后：", joined)
        self.assertIn("任务状态变为「处理中」。", joined)
        self.assertIn("注意：", joined)
        self.assertIn("处理中不要重复提交。", joined)

    def test_legacy_table_section(self) -> None:
        data = self.load_fixture("legacy.json")
        data["modules"][0]["sections"].append(
            {
                "heading": "额度说明",
                "kind": "table",
                "rows": [["免费额度", "每月 100 次"], ["超出后", "需要联系管理员"]],
            }
        )
        texts = all_texts(self.open_docx(self.render_data(data)))
        joined = "\n".join(texts)
        self.assertIn("额度说明", joined)
        self.assertIn("免费额度", joined)
        self.assertIn("每月 100 次", joined)

    def test_table_section_dict_rows(self) -> None:
        data = self.load_fixture("legacy.json")
        data["modules"][0]["sections"].append(
            {
                "heading": "账号规则",
                "kind": "table",
                "rows": [{"key": "登录方式", "value": "企业账号"}, {"项目": "密码策略", "说明": "每 90 天更换"}],
            }
        )
        texts = all_texts(self.open_docx(self.render_data(data)))
        joined = "\n".join(texts)
        self.assertIn("登录方式", joined)
        self.assertIn("企业账号", joined)
        self.assertIn("密码策略", joined)
        self.assertIn("每 90 天更换", joined)

    def test_legacy_section_empty_content_no_heading(self) -> None:
        data = self.load_fixture("legacy.json")
        data["modules"][0]["sections"].append({"heading": "空内容标题", "kind": "steps", "items": []})
        data["modules"][0]["sections"].append({"heading": "空表格标题", "kind": "table", "rows": []})
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertNotIn("空内容标题", paras)
        self.assertNotIn("空表格标题", paras)

    def test_legacy_empty_string_steps_do_not_break_numbering(self) -> None:
        data = self.load_fixture("legacy.json")
        data["modules"][1]["steps"] = ["", "选择趋势方向。", "", "点击「开始生成」。"]
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertIn("1. 选择趋势方向。", paras)
        self.assertIn("2. 点击「开始生成」。", paras)
        self.assertNotIn("3. 点击「开始生成」。", paras)

    def test_legacy_missing_image_is_safe(self) -> None:
        out = self.render_fixture("legacy.json")
        doc = self.open_docx(out)
        paras = [p.text for p in doc.paragraphs]
        self.assertTrue(any("截图缺失" in p for p in paras))

    def test_legacy_cli(self) -> None:
        output = self.tmpdir / "cli-legacy.docx"
        result = subprocess.run(
            [sys.executable, str(SCRIPT), "--input", str(FIXTURES / "legacy.json"), "--output", str(output)],
            capture_output=True,
            text=True,
        )
        self.assertEqual(result.returncode, 0, result.stderr)
        self.assertTrue(output.exists())
        self.open_docx(output)

    def test_v2_cli(self) -> None:
        output = self.tmpdir / "cli-full-v2.docx"
        result = subprocess.run(
            [sys.executable, str(SCRIPT), "--input", str(FIXTURES / "full-v2.json"), "--output", str(output)],
            capture_output=True,
            text=True,
        )
        self.assertEqual(result.returncode, 0, result.stderr)
        self.assertTrue(output.exists())
        paras = [p.text for p in self.open_docx(output).paragraphs]
        self.assertIn("常用任务", paras)


class TestV2(RendererTestBase):
    def test_at_a_glance_first_page(self) -> None:
        out = self.render_fixture("full-v2.json")
        paras = [p.text for p in self.open_docx(out).paragraphs]
        self.assertIn("3 分钟了解这个产品", paras)
        self.assertIn("这个产品能帮你做什么", paras)
        self.assertIn("最常用的几件事", paras)
        self.assertIn("第一次建议这样用", paras)
        self.assertIn("开始前要知道", paras)
        self.assertIn("系统入口", paras)
        self.assertTrue(any("帮助同事快速了解最近值得关注的内容方向" in p for p in paras))

    def test_tasks_priority_over_modules(self) -> None:
        out = self.render_fixture("full-v2.json")
        paras = [p.text for p in self.open_docx(out).paragraphs]
        self.assertLess(paras.index("常用任务"), paras.index("功能索引"))
        self.assertNotIn("模块使用说明", paras)
        # Index mode: module sections are not re-rendered (no duplication).
        self.assertNotIn("怎么操作", paras)
        self.assertIn("查看热点与内容方向。", paras)

    def test_object_step_callouts(self) -> None:
        out = self.render_fixture("full-v2.json")
        texts = all_texts(self.open_docx(out))
        joined = "\n".join(texts)
        self.assertIn("1. 点击左侧菜单「趋势洞察」。", joined)
        self.assertIn("完成后：", joined)
        self.assertIn("页面显示热点排行和时间筛选。", joined)
        self.assertIn("注意：", joined)
        self.assertIn("页面还在加载时不要重复点击。", joined)
        self.assertIn("建议：", joined)
        self.assertIn("可以先按时间范围缩小列表。", joined)

    def test_string_step_and_expected_result(self) -> None:
        out = self.render_fixture("full-v2.json")
        texts = all_texts(self.open_docx(out))
        joined = "\n".join(texts)
        self.assertIn("2. 点击「开始生成」。", joined)
        self.assertIn("任务状态变为「处理中」", joined)

    def test_step_screenshot_local_dedupe_and_cross_section_repeat(self) -> None:
        out = self.render_fixture("full-v2.json")
        doc = self.open_docx(out)
        # Same image referenced at step + task level inside task 1 is
        # deduplicated; the same file may repeat in other tasks.
        self.assertEqual(len(doc.inline_shapes), 2)
        paras = [p.text for p in doc.paragraphs]
        self.assertTrue(any("图 1：趋势列表" in p for p in paras))
        self.assertTrue(any("图 2：历史记录页面" in p for p in paras))

    def test_common_problems(self) -> None:
        out = self.render_fixture("full-v2.json")
        paras = [p.text for p in self.open_docx(out).paragraphs]
        self.assertIn("Q：为什么这里没有可以选择的内容？", paras)
        self.assertIn("请确认账号权限或联系管理员开通。", paras)

    def test_priority_tags(self) -> None:
        out = self.render_fixture("full-v2.json")
        paras = [p.text for p in self.open_docx(out).paragraphs]
        self.assertTrue(any("（核心）" in p for p in paras))
        self.assertTrue(any("（辅助）" in p for p in paras))
        self.assertTrue(any("（参考）" in p for p in paras))

    def test_core_result_section(self) -> None:
        out = self.render_fixture("full-v2.json")
        paras = [p.text for p in self.open_docx(out).paragraphs]
        self.assertIn("完成后你会看到", paras)

    def test_minimal_v2(self) -> None:
        out = self.render_fixture("minimal-v2.json")
        doc = self.open_docx(out)
        paras = [p.text for p in doc.paragraphs]
        self.assertIn("3 分钟了解这个产品", paras)
        self.assertIn("常用任务", paras)
        self.assertIn("1. 打开系统并登录。", paras)
        self.assertNotIn("推荐使用流程", paras)
        self.assertNotIn("模块使用说明", paras)
        self.assertNotIn("问题反馈模板", paras)

    def test_unknown_priority_defaults_to_core(self) -> None:
        data = self.load_fixture("minimal-v2.json")
        data["tasks"][0]["priority"] = "high"
        data["tasks"][0]["result"] = "任务完成后结果保存在「历史记录」"
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertIn("完成后你会看到", paras)
        self.assertIn("任务完成后结果保存在「历史记录」", paras)

    def test_missing_action_step_gets_placeholder(self) -> None:
        data = self.load_fixture("minimal-v2.json")
        data["tasks"][0]["steps"] = [
            {"expected_result": "该步骤缺少操作说明，但结果仍应展示。"},
            "正常的第二步。",
        ]
        texts = all_texts(self.open_docx(self.render_data(data)))
        joined = "\n".join(texts)
        self.assertIn("1. （此步骤缺少操作说明）", joined)
        self.assertIn("完成后：", joined)
        self.assertIn("该步骤缺少操作说明，但结果仍应展示。", joined)
        self.assertIn("2. 正常的第二步。", joined)

    def test_action_only_step_creates_no_empty_callouts(self) -> None:
        data = self.load_fixture("minimal-v2.json")
        data["tasks"][0]["steps"] = [{"action": "仅操作，无任何补充说明。"}]
        texts = all_texts(self.open_docx(self.render_data(data)))
        joined = "\n".join(texts)
        self.assertIn("1. 仅操作，无任何补充说明。", joined)
        self.assertNotIn("完成后：", joined)
        self.assertNotIn("注意：", joined)
        self.assertNotIn("建议：", joined)

    def test_task_level_screenshot_fallback(self) -> None:
        data = self.load_fixture("minimal-v2.json")
        data["tasks"][0]["screenshots"] = [
            {"path": str(self.png), "caption": "任务级兜底截图"}
        ]
        doc = self.open_docx(self.render_data(data))
        self.assertEqual(len(doc.inline_shapes), 1)
        paras = [p.text for p in doc.paragraphs]
        self.assertTrue(any("图 1：任务级兜底截图" in p for p in paras))

    def test_empty_steps_or_result_do_not_create_headings(self) -> None:
        data = self.load_fixture("minimal-v2.json")
        data["tasks"][0]["steps"] = ""
        data["tasks"][0]["result"] = ""
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertNotIn("操作步骤", paras)
        self.assertNotIn("完成后你会看到", paras)

    def test_empty_string_steps_do_not_break_numbering(self) -> None:
        data = self.load_fixture("minimal-v2.json")
        data["tasks"][0]["steps"] = ["", "第一步", "", "第二步"]
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertIn("1. 第一步", paras)
        self.assertIn("2. 第二步", paras)
        self.assertFalse(any(p.startswith("3. ") for p in paras))


class TestBoundaries(RendererTestBase):
    def test_empty_quick_start_no_section(self) -> None:
        data = self.load_fixture("full-v2.json")
        data["quick_start"] = []
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertNotIn("第一次建议这样用", paras)
        self.assertNotIn("快速开始", paras)

    def test_empty_recommended_workflow_no_section(self) -> None:
        data = self.load_fixture("full-v2.json")
        data["recommended_workflow"] = []
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertNotIn("推荐使用流程", paras)

    def test_recommended_workflow_duplicate_skipped(self) -> None:
        data = self.load_fixture("full-v2.json")
        data["recommended_workflow"] = list(data["quick_start"])
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertNotIn("推荐使用流程", paras)

    def test_recommended_workflow_prefix_skipped(self) -> None:
        data = self.load_fixture("full-v2.json")
        data["recommended_workflow"] = data["quick_start"][:2]
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertNotIn("推荐使用流程", paras)

    def test_recommended_workflow_extending_quick_start_merged_once(self) -> None:
        data = self.load_fixture("full-v2.json")
        data["quick_start"] = ["第一步。", "第二步。"]
        data["recommended_workflow"] = ["第一步。", "第二步。", "第三步（推荐流程独有）。"]
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertNotIn("推荐使用流程", paras)
        self.assertIn("3. 第三步（推荐流程独有）。", paras)

    def test_top_tasks_capped_at_five(self) -> None:
        data = self.load_fixture("minimal-v2.json")
        data["at_a_glance"]["top_tasks"] = ["任务一", "任务二", "任务三", "任务四", "任务五", "任务六", "任务七"]
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertIn("任务五", paras)
        self.assertNotIn("任务六", paras)

    def test_overview_not_duplicated_when_identical(self) -> None:
        data = self.load_fixture("full-v2.json")
        same = data["at_a_glance"]["what_it_does"]
        data["overview"] = same
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertEqual(sum(1 for p in paras if p == same), 1)
        self.assertNotIn("简短说明", paras)

    def test_overview_extra_info_rendered_standalone(self) -> None:
        data = self.load_fixture("full-v2.json")
        data["overview"] = "更详细的补充说明：包含使用范围与适用场景。"
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertIn("简短说明", paras)
        self.assertIn("更详细的补充说明：包含使用范围与适用场景。", paras)

    def test_usage_rules_not_duplicated_between_page1_and_section(self) -> None:
        # Legacy: rules are shown on page 1 only (fallback), never twice.
        out = self.render_fixture("legacy.json")
        paras = [p.text for p in self.open_docx(out).paragraphs]
        self.assertEqual(sum(1 for p in paras if p == "账号仅限内部使用。"), 1)
        self.assertNotIn("使用规范", paras)
        # V2 with distinct rules: section renders them once.
        out2 = self.render_fixture("full-v2.json")
        paras2 = [p.text for p in self.open_docx(out2).paragraphs]
        self.assertEqual(sum(1 for p in paras2 if p == "账号仅限内部使用。"), 1)
        self.assertIn("使用规范", paras2)

    def test_usage_rules_remainder_rendered_when_more_than_page1(self) -> None:
        data = self.load_fixture("minimal-v2.json")
        data["usage_rules"] = ["规则一", "规则二", "规则三", "规则四", "规则五"]
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertIn("使用规范", paras)
        self.assertEqual(sum(1 for p in paras if p == "规则五"), 1)

    def test_default_title_header_has_no_broken_product_name(self) -> None:
        doc = self.open_docx(self.render_data({"version_date": "2026-08-12"}))
        header_text = " ".join(p.text for p in doc.sections[0].header.paragraphs)
        self.assertEqual(header_text, "公司内部使用指南")

    def test_corrupted_image_error_deduped(self) -> None:
        bad = self.tmpdir / "bad.png"
        bad.write_bytes(b"this is not a png")
        data = {
            "title": "损坏图片测试",
            "tasks": [
                {
                    "title": "测试任务",
                    "priority": "core",
                    "steps": [
                        {
                            "action": "第一步",
                            "screenshot": {"path": str(bad), "caption": "损坏图片"},
                        }
                    ],
                    "screenshots": [{"path": str(bad), "caption": "损坏图片"}],
                }
            ],
        }
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        errors = [p for p in paras if p.startswith("截图无法插入")]
        self.assertEqual(len(errors), 1)

    def test_figure_counter_skips_failed_image(self) -> None:
        bad = self.tmpdir / "bad2.png"
        bad.write_bytes(b"not a png either")
        data = {
            "title": "图号连续性测试",
            "tasks": [
                {
                    "title": "测试任务",
                    "priority": "core",
                    "steps": [
                        {"action": "失败步骤", "screenshot": {"path": str(bad), "caption": "损坏图片"}},
                        {"action": "成功步骤", "screenshot": {"path": str(self.png), "caption": "正常图片"}},
                    ],
                }
            ],
        }
        doc = self.open_docx(self.render_data(data))
        self.assertEqual(len(doc.inline_shapes), 1)
        paras = [p.text for p in doc.paragraphs]
        self.assertTrue(any("图 1：正常图片" in p for p in paras))

    def test_uncovered_module_renders_full_under_index(self) -> None:
        data = self.load_fixture("full-v2.json")
        data["modules"].append(
            {
                "name": "系统设置",
                "purpose": "管理账号与偏好设置。",
                "sections": [
                    {"heading": "怎么操作", "kind": "steps", "items": ["进入「系统设置」。", "修改偏好后保存。"]}
                ],
            }
        )
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertIn("功能索引", paras)
        # Covered modules stay as index entries (their sections are not
        # re-rendered); the uncovered module renders its section exactly once.
        self.assertEqual(sum(1 for p in paras if p == "怎么操作"), 1)
        self.assertIn("1. 进入「系统设置」。", paras)

    def test_module_substring_not_misclassified_as_covered(self) -> None:
        data = self.load_fixture("full-v2.json")
        data["modules"].append(
            {
                "name": "记录",
                "purpose": "通用记录查询。",
                "sections": [{"heading": "使用说明", "kind": "bullets", "items": ["按条件筛选记录。"]}],
            }
        )
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        # Module「历史记录」is covered (entry「历史记录」); module「记录」is
        # not an exact match and must render in full.
        self.assertIn("使用说明", paras)
        self.assertIn("按条件筛选记录。", paras)

    def test_faq_string_entries(self) -> None:
        data = self.load_fixture("minimal-v2.json")
        data["faq"] = ["为什么登录失败？", "为什么页面打不开？"]
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertIn("为什么登录失败？", paras)
        self.assertIn("为什么页面打不开？", paras)

    def test_faq_answer_only_entry_is_kept(self) -> None:
        data = self.load_fixture("minimal-v2.json")
        data["faq"] = [{"answer": "只有回答、没有问题的条目仍应显示。"}]
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertIn("只有回答、没有问题的条目仍应显示。", paras)

    def test_common_problems_without_questions_no_heading(self) -> None:
        data = self.load_fixture("minimal-v2.json")
        data["tasks"][0]["common_problems"] = [{"answer": "只有回答。"}]
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertNotIn("常见问题", paras)

    def test_structured_values_never_leak_python_repr(self) -> None:
        data = {
            "title": "无泄漏测试",
            "at_a_glance": {"what_it_does": ["帮助", "协作"]},
            "tasks": [{"title": {"not": "a string"}, "priority": "core", "steps": ["正常步骤。"]}],
        }
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertIn("帮助；协作", paras)
        self.assertTrue(all("{" not in p and "[" not in p and "'" not in p for p in paras if p))

    def test_reference_task_does_not_render_step_screenshots(self) -> None:
        data = self.load_fixture("minimal-v2.json")
        data["tasks"][0]["priority"] = "reference"
        data["tasks"][0]["steps"] = [
            {"action": "打开「账号设置」。", "screenshot": {"path": str(self.png), "caption": "设置页"}}
        ]
        doc = self.open_docx(self.render_data(data))
        self.assertEqual(len(doc.inline_shapes), 0)
        paras = [p.text for p in doc.paragraphs]
        self.assertIn("1. 打开「账号设置」。", paras)
        self.assertFalse(any("设置页" in p for p in paras))

    def test_legacy_module_empty_dict_fields_no_empty_headings(self) -> None:
        data = self.load_fixture("legacy.json")
        data["modules"][1]["scenarios"] = {}
        data["modules"][1]["key_outputs"] = {}
        data["modules"][1]["notes"] = {}
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertNotIn("适用场景", paras)
        self.assertNotIn("重点结果", paras)
        self.assertNotIn("注意事项", paras)

    def test_header_product_name_suffix_extraction(self) -> None:
        doc = self.open_docx(self.render_data({"title": "趋势分析平台 使用说明书"}))
        header_text = " ".join(p.text for p in doc.sections[0].header.paragraphs)
        self.assertEqual(header_text, "趋势分析平台 · 公司内部使用指南")

    def test_numbered_blocks_restart_at_one(self) -> None:
        out = self.render_fixture("full-v2.json")
        paras = [p.text for p in self.open_docx(out).paragraphs if p.text.strip()]
        blocks = []
        current: list[int] = []
        for p in paras:
            match = re.match(r"^(\d+)\. ", p)
            if match:
                current.append(int(match.group(1)))
            elif current:
                blocks.append(current)
                current = []
        if current:
            blocks.append(current)
        self.assertGreaterEqual(len(blocks), 4)
        for block in blocks:
            self.assertEqual(block, list(range(1, len(block) + 1)))

    def test_no_problem_feedback_template(self) -> None:
        for name in ("legacy.json", "minimal-v2.json", "full-v2.json"):
            out = self.render_fixture(name)
            paras = [p.text for p in self.open_docx(out).paragraphs]
            self.assertNotIn("问题反馈模板", paras)

    def test_all_missing_images_do_not_crash(self) -> None:
        data = self.load_fixture("full-v2.json")
        data = self.replace_image_paths(data, "images/sample.png", "images/does-not-exist.png")
        out = self.render_data(data)
        paras = [p.text for p in self.open_docx(out).paragraphs]
        self.assertTrue(any("截图缺失" in p for p in paras))

    def test_malformed_screenshot_entries_do_not_crash(self) -> None:
        data = self.load_fixture("minimal-v2.json")
        data["tasks"][0]["screenshots"] = [123, None, ["nested"], {"path": str(self.png), "caption": "正常截图"}]
        doc = self.open_docx(self.render_data(data))
        self.assertEqual(len(doc.inline_shapes), 1)
        paras = [p.text for p in doc.paragraphs]
        self.assertTrue(any("图 1：正常截图" in p for p in paras))

    def test_legacy_numbered_section_kind_alias(self) -> None:
        data = self.load_fixture("legacy.json")
        data["modules"][0]["sections"].append(
            {"heading": "操作顺序", "kind": "numbered", "items": ["先做甲。", "再做乙。"]}
        )
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertIn("1. 先做甲。", paras)
        self.assertIn("2. 再做乙。", paras)

    def test_empty_document(self) -> None:
        out = self.render_data({}, "empty.docx")
        doc = self.open_docx(out)
        self.assertEqual(heading_texts(doc), [])
        self.assertEqual(len(doc.tables), 0)

    def test_minimal_v2_has_no_metadata_table(self) -> None:
        out = self.render_fixture("minimal-v2.json")
        doc = self.open_docx(out)
        self.assertEqual(len(doc.tables), 0)

    def test_before_you_start_capped_at_four(self) -> None:
        data = self.load_fixture("minimal-v2.json")
        data["at_a_glance"]["before_you_start"] = ["注意一", "注意二", "注意三", "注意四", "注意五", "注意六"]
        paras = [p.text for p in self.open_docx(self.render_data(data)).paragraphs]
        self.assertIn("注意四", paras)
        self.assertNotIn("注意五", paras)

    def test_no_empty_headings(self) -> None:
        data = {
            "title": "测试文档",
            "at_a_glance": {},
            "tasks": [{"title": ""}],
            "modules": [],
            "quick_start": [],
            "recommended_workflow": [],
            "faq": [],
        }
        doc = self.open_docx(self.render_data(data))
        self.assertTrue(all(h for h in heading_texts(doc)))

    def test_header_and_footer(self) -> None:
        out = self.render_fixture("full-v2.json")
        doc = self.open_docx(out)
        header_text = " ".join(p.text for p in doc.sections[0].header.paragraphs)
        self.assertIn("内容趋势分析平台 · 公司内部使用指南", header_text)
        footer_xml = doc.sections[0].footer.paragraphs[0]._p.xml
        self.assertIn("PAGE", footer_xml)
        self.assertIn("版本日期：2026-08-12", footer_xml)

    def test_legacy_quick_start_on_first_page(self) -> None:
        out = self.render_fixture("legacy.json")
        paras = [p.text for p in self.open_docx(out).paragraphs]
        self.assertIn("第一次建议这样用", paras)
        self.assertIn("1. 打开系统并登录。", paras)

    def test_legacy_overview_rendered_once(self) -> None:
        out = self.render_fixture("legacy.json")
        paras = [p.text for p in self.open_docx(out).paragraphs]
        overview_text = "内容趋势分析平台用于帮助同事快速了解热点方向，并基于热点生成自己的内容。"
        self.assertEqual(sum(1 for p in paras if p == overview_text), 1)
        self.assertNotIn("简短说明", paras)

    def test_recommended_workflow_distinct_rendered(self) -> None:
        out = self.render_fixture("full-v2.json")
        paras = [p.text for p in self.open_docx(out).paragraphs]
        self.assertIn("推荐使用流程", paras)
        self.assertIn("1. 先看趋势确认方向。", paras)


class TestV21Fixes(RendererTestBase):
    """Regression tests for the V2.1 hardening round."""

    def module_data(self) -> dict:
        return {
            "name": "趋势洞察",
            "purpose": "查看热点与内容方向。",
            "sections": [
                {"heading": "怎么操作", "kind": "steps", "items": ["打开「趋势洞察」。", "点击「查看」。"]}
            ],
        }

    def render_with(self, tasks, modules, at_a_glance=None, **extra):
        data = {"title": "V2.1 测试", "tasks": tasks, "modules": modules}
        if at_a_glance is not None:
            data["at_a_glance"] = at_a_glance
        data.update(extra)
        return self.render_data(data)

    def test_covers_modules_explicit_mapping(self) -> None:
        tasks = [
            {
                "title": "我要看看最近有什么热点",
                "entry": "左侧菜单 > 趋势洞察",
                "covers_modules": ["趋势洞察"],
                "priority": "core",
                "steps": ["打开「趋势洞察」。", "点击「查看」。"],
            }
        ]
        paras = [p.text for p in self.open_docx(self.render_with(tasks, [self.module_data()])).paragraphs]
        self.assertIn("功能索引", paras)
        self.assertIn("查看热点与内容方向。", paras)
        # Covered module stays an index entry: sections are not re-rendered.
        self.assertNotIn("怎么操作", paras)

    def test_covers_modules_overrides_entry_heuristic(self) -> None:
        tasks = [
            {
                "title": "我要看看最近有什么热点",
                "entry": "通过「趋势洞察」进入",
                "covers_modules": ["另一个模块"],
                "priority": "core",
                "steps": ["打开系统。"],
            }
        ]
        paras = [p.text for p in self.open_docx(self.render_with(tasks, [self.module_data()])).paragraphs]
        # Explicit covers_modules wins: 「趋势洞察」 in entry must NOT mark
        # the module as covered; it renders in full.
        self.assertIn("怎么操作", paras)
        self.assertIn("1. 打开「趋势洞察」。", paras)

    def test_heuristic_fallback_without_covers_modules(self) -> None:
        tasks = [
            {
                "title": "我要看看最近有什么热点",
                "entry": "左侧菜单「趋势洞察」",
                "priority": "core",
                "steps": ["打开「趋势洞察」。", "点击「查看」。"],
            }
        ]
        paras = [p.text for p in self.open_docx(self.render_with(tasks, [self.module_data()])).paragraphs]
        self.assertIn("功能索引", paras)
        self.assertNotIn("怎么操作", paras)

    def test_supporting_result_rendered(self) -> None:
        tasks = [
            {
                "title": "我要下载历史结果",
                "priority": "supporting",
                "steps": ["打开「历史记录」。", "点击「下载」。"],
                "result": "点击下载后浏览器开始保存文件。",
            }
        ]
        paras = [p.text for p in self.open_docx(self.render_with(tasks, [])).paragraphs]
        self.assertIn("完成后你会看到", paras)
        self.assertIn("点击下载后浏览器开始保存文件。", paras)

    def test_supporting_common_problems_rendered(self) -> None:
        tasks = [
            {
                "title": "我要下载历史结果",
                "priority": "supporting",
                "steps": ["打开「历史记录」。", "点击「下载」。"],
                "common_problems": [
                    {"question": "为什么下载没有反应？", "answer": "请确认任务已经完成。"}
                ],
            }
        ]
        paras = [p.text for p in self.open_docx(self.render_with(tasks, [])).paragraphs]
        self.assertIn("Q：为什么下载没有反应？", paras)
        self.assertIn("请确认任务已经完成。", paras)

    def test_top_tasks_fallback_from_tasks(self) -> None:
        tasks = [
            {"title": "我要查看最近热点", "priority": "core", "steps": ["打开系统。"]},
            {"title": "我要生成内容", "priority": "core", "steps": ["打开系统。"]},
        ]
        at_a_glance = {"what_it_does": "一句话说明。"}
        paras = [p.text for p in self.open_docx(self.render_with(tasks, [], at_a_glance)).paragraphs]
        self.assertIn("最常用的几件事", paras)
        self.assertIn("我要查看最近热点", paras)
        self.assertIn("我要生成内容", paras)

    def test_top_tasks_primary_source_deduped(self) -> None:
        tasks = [{"title": "我要查看最近热点", "priority": "core", "steps": ["打开系统。"]}]
        at_a_glance = {"top_tasks": ["热点", "热点", "内容", "内容", "历史", "历史"]}
        paras = [p.text for p in self.open_docx(self.render_with(tasks, [], at_a_glance)).paragraphs]
        self.assertEqual(sum(1 for p in paras if p == "热点"), 1)
        self.assertEqual(sum(1 for p in paras if p == "内容"), 1)
        self.assertIn("历史", paras)

    def test_tasks_fallback_priority_over_modules(self) -> None:
        tasks = [
            {"title": "我要查看最近热点", "priority": "core", "steps": ["打开系统。"]},
            {"title": "我要生成内容", "priority": "core", "steps": ["打开系统。"]},
        ]
        modules = [{"name": "模块甲", "purpose": "甲。"}, {"name": "模块乙", "purpose": "乙。"}]
        paras = [p.text for p in self.open_docx(self.render_with(tasks, modules)).paragraphs]
        # Page 1 uses task titles; module names only appear later in the index.
        self.assertLess(paras.index("我要查看最近热点"), paras.index("模块甲"))
        self.assertIn("我要生成内容", paras)

    def test_common_problem_answer_list_no_repr(self) -> None:
        tasks = [
            {
                "title": "生成失败处理",
                "priority": "core",
                "steps": ["打开系统。"],
                "common_problems": [{"question": "生成失败怎么办？", "answer": ["稍后重试", "联系管理员"]}],
            }
        ]
        paras = [p.text for p in self.open_docx(self.render_with(tasks, [])).paragraphs]
        self.assertIn("稍后重试；联系管理员", paras)
        self.assertFalse(any("[" in p and "]" in p for p in paras))

    def test_dict_never_leaks_repr(self) -> None:
        tasks = [
            {
                "title": "我要查看热点",
                "priority": "core",
                "purpose": {"unexpected": "value"},
                "steps": ["打开系统。"],
            }
        ]
        paras = [p.text for p in self.open_docx(self.render_with(tasks, [])).paragraphs]
        self.assertNotIn("unexpected", paras)
        self.assertFalse(any("{" in p or "'" in p for p in paras))

    def test_same_task_screenshot_deduped(self) -> None:
        tasks = [
            {
                "title": "趋势任务",
                "priority": "core",
                "steps": [
                    {"action": "点击「查看」。", "screenshot": {"path": str(self.png), "caption": "步骤截图"}}
                ],
                "screenshots": [{"path": str(self.png), "caption": "任务兜底截图"}],
            }
        ]
        doc = self.open_docx(self.render_with(tasks, []))
        self.assertEqual(len(doc.inline_shapes), 1)
        paras = [p.text for p in doc.paragraphs]
        self.assertTrue(any("图 1：步骤截图" in p for p in paras))

    def test_cross_task_same_image_repeated(self) -> None:
        tasks = [
            {
                "title": "任务甲",
                "priority": "core",
                "steps": [{"action": "打开首页。", "screenshot": {"path": str(self.png), "caption": "首页"}}],
            },
            {
                "title": "任务乙",
                "priority": "core",
                "steps": [{"action": "回到首页。", "screenshot": {"path": str(self.png), "caption": "首页"}}],
            },
        ]
        doc = self.open_docx(self.render_with(tasks, []))
        self.assertEqual(len(doc.inline_shapes), 2)
        paras = [p.text for p in doc.paragraphs]
        self.assertTrue(any("图 1：首页" in p for p in paras))
        self.assertTrue(any("图 2：首页" in p for p in paras))


if __name__ == "__main__":
    unittest.main()
