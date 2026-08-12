#!/usr/bin/env python3
"""Render a structured product usage guide JSON to a colleague-friendly DOCX.

Supports both the legacy schema (overview / quick_start / modules / ...) and
the V2 task-oriented schema (at_a_glance / tasks / object steps / ...).
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ---------------------------------------------------------------------------
# Small helpers
# ---------------------------------------------------------------------------


def text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def as_list(value: Any) -> list[Any]:
    if value is None or value == "":
        return []
    if isinstance(value, list):
        return value
    return [value]


def plain_text(value: Any) -> str:
    """Unified user-facing text helper.

    - str: trimmed as-is.
    - list: string items joined with 「；」; non-string items are skipped.
    - anything else (dict, number, bool, nested structures): empty string.

    Never leaks a Python repr into the document.
    """
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, list):
        return "；".join(item.strip() for item in value if isinstance(item, str) and item.strip())
    return ""


def set_run_font(run: Any, east_asia: str = "微软雅黑") -> None:
    run.font.name = east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def prevent_row_split(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def normalize_path_key(path: Path) -> str:
    return os.path.normcase(str(path.resolve()))


def product_name_from_title(title: str) -> str:
    stripped_title = title.strip()
    if stripped_title in ("产品使用说明总结", "产品使用指南", "产品使用说明"):
        return ""
    for suffix in ("公司内部同事使用说明", "公司内部使用说明", "使用说明书", "使用说明", "说明书", "使用指南"):
        if suffix in title:
            return title.split(suffix, 1)[0].strip()
    return stripped_title


class DocState:
    """Per-document mutable state (global figure numbering)."""

    def __init__(self) -> None:
        self.figure_counter = 0


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    for style_name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    normal = styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(4)

    h1 = styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.size = Pt(13)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.size = Pt(11.5)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    styles["Title"].font.size = Pt(20)
    styles["Title"].paragraph_format.space_after = Pt(6)
    styles["Title"].font.color.rgb = RGBColor(0, 0, 0)
    styles["Subtitle"].font.size = Pt(10.5)
    styles["Subtitle"].paragraph_format.space_after = Pt(8)
    styles["Subtitle"].font.color.rgb = RGBColor(0x59, 0x59, 0x59)


def add_header_footer(doc: Document, title: str, version_date: str) -> None:
    product = product_name_from_title(title)
    header_text = f"{product} · 公司内部使用指南" if product else "公司内部使用指南"

    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = ""
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run(header_text)
    header_run.font.size = Pt(9)
    set_run_font(header_run)

    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = ""
    footer_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

    left_run = footer_para.add_run(f"版本日期：{plain_text(version_date) or '未注明'} | ")
    left_run.font.size = Pt(9)
    set_run_font(left_run)
    footer_para.add_run("\t")
    pre_run = footer_para.add_run("第 ")
    pre_run.font.size = Pt(9)
    set_run_font(pre_run)
    _add_page_field(footer_para)
    post_run = footer_para.add_run(" 页")
    post_run.font.size = Pt(9)
    set_run_font(post_run)


def _add_page_field(paragraph: Any) -> None:
    begin_run = paragraph.add_run()
    begin_run.font.size = Pt(9)
    set_run_font(begin_run)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(fld_begin)

    instr_run = paragraph.add_run()
    instr_run.font.size = Pt(9)
    set_run_font(instr_run)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    instr_run._r.append(instr)

    # separate + cached value so the field has a display result before Word
    # refreshes it.
    sep_run = paragraph.add_run()
    sep_run.font.size = Pt(9)
    set_run_font(sep_run)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(fld_sep)
    sep_run.add_text("1")

    end_run = paragraph.add_run()
    end_run.font.size = Pt(9)
    set_run_font(end_run)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)


# ---------------------------------------------------------------------------
# Basic content blocks
# ---------------------------------------------------------------------------


def add_paragraph(doc: Document, value: Any, style: str | None = None) -> None:
    content = plain_text(value)
    if not content:
        return
    para = doc.add_paragraph(style=style)
    run = para.add_run(content)
    set_run_font(run)


def add_bullets(doc: Document, items: Any) -> None:
    for item in as_list(items):
        content = plain_text(item)
        if not content:
            continue
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(content)
        set_run_font(run)


def add_steps(
    doc: Document,
    items: Any,
    base_dir: Path,
    state: DocState,
    seen: set[str],
    with_screenshots: bool = True,
) -> None:
    """Numbered steps. Accepts plain strings and V2 step objects.

    Every independent step block restarts at ``1.``. Step objects may carry
    ``expected_result`` / ``warning`` / ``tip`` callouts and a ``screenshot``.
    """
    items = [item for item in as_list(items) if isinstance(item, dict) or plain_text(item)]
    for index, item in enumerate(items, start=1):
        if isinstance(item, dict):
            action = plain_text(item.get("action")) or plain_text(item.get("text"))
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.keep_with_next = True
            run = para.add_run(f"{index}. {action}" if action else f"{index}. （此步骤缺少操作说明）")
            set_run_font(run)

            expected = plain_text(item.get("expected_result"))
            warning = plain_text(item.get("warning"))
            tip = plain_text(item.get("tip"))
            if expected:
                add_callout(doc, "完成后", expected, "E8F1E8")
            if warning:
                add_callout(doc, "注意", warning, "FDF0E6")
            if tip:
                add_callout(doc, "建议", tip, "EAF1F8")
            if with_screenshots and item.get("screenshot"):
                add_screenshots(doc, item.get("screenshot"), base_dir, state, seen)
        else:
            content = plain_text(item)
            if not content:
                continue
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.keep_with_next = True
            run = para.add_run(f"{index}. {content}")
            set_run_font(run)


def clean_steps(items: Any) -> list[Any]:
    """Drop empty string steps before rendering so numbering stays 1..n."""
    cleaned = []
    for item in as_list(items):
        if isinstance(item, dict):
            cleaned.append(item)
        elif plain_text(item):
            cleaned.append(item)
    return cleaned


def add_callout(doc: Document, label: str, content: Any, fill: str) -> None:
    body = plain_text(content)
    if not body:
        return
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_shading(cell, fill)
    prevent_row_split(table.rows[0])
    label_para = cell.paragraphs[0]
    label_para.paragraph_format.space_before = Pt(2)
    label_para.paragraph_format.space_after = Pt(1)
    label_run = label_para.add_run(f"{label}：")
    label_run.bold = True
    set_run_font(label_run)
    body_para = cell.add_paragraph()
    body_para.paragraph_format.space_before = Pt(0)
    body_para.paragraph_format.space_after = Pt(2)
    body_run = body_para.add_run(body)
    set_run_font(body_run)

    # Word requires a paragraph between adjacent tables; keep it invisible.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)
    spacer_run = spacer.add_run("")
    spacer_run.font.size = Pt(2)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    rows = [(k, v) for k, v in rows if plain_text(v)]
    if not rows:
        return
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "项目"
    header[1].text = "说明"
    for cell in header:
        set_cell_shading(cell, "D9EAF7")
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value


def resolve_path(raw_path: str, base_dir: Path) -> Path:
    path = Path(raw_path)
    if path.is_absolute():
        return path
    return base_dir / path


def add_screenshots(
    doc: Document,
    screenshots: Any,
    base_dir: Path,
    state: DocState,
    seen: set[str],
    image_width: float = 5.8,
) -> None:
    for shot in as_list(screenshots):
        if isinstance(shot, str):
            path_value = shot
            caption = ""
        elif isinstance(shot, dict):
            path_value = plain_text(shot.get("path"))
            caption = plain_text(shot.get("caption"))
        else:
            # Malformed entry (number, null, nested list): never crash.
            continue
        if not path_value:
            continue
        image_path = resolve_path(path_value, base_dir)
        key = normalize_path_key(image_path)
        if key in seen:
            continue
        if not image_path.exists():
            add_paragraph(doc, f"截图缺失：{path_value}")
            seen.add(key)
            continue
        para = None
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.keep_with_next = True
            run = para.add_run()
            run.add_picture(str(image_path), width=Inches(image_width))
            state.figure_counter += 1
            seen.add(key)

            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(8)
            caption_text = f"图 {state.figure_counter}" + (f"：{caption}" if caption else "")
            cap_run = cap.add_run(caption_text)
            cap_run.italic = True
            cap_run.font.size = Pt(9)
            set_run_font(cap_run)
        except Exception as exc:
            seen.add(key)
            if para is not None:
                para._element.getparent().remove(para._element)
            print(f"截图无法插入：{path_value}（{exc}）", file=sys.stderr)
            add_paragraph(doc, f"截图无法插入：{path_value}（图片无法识别或已损坏）")


# ---------------------------------------------------------------------------
# Legacy sections / modules
# ---------------------------------------------------------------------------


def add_section(
    doc: Document,
    section: dict[str, Any],
    base_dir: Path,
    state: DocState,
    seen: set[str],
) -> None:
    heading = plain_text(section.get("heading"))
    kind = text(section.get("kind")).lower() or "paragraph"
    if kind in {"steps", "numbered", "numbers"}:
        items = clean_steps(section.get("items") or section.get("body"))
        if not items:
            return
        if heading:
            doc.add_heading(heading, level=3)
        add_steps(doc, items, base_dir, state, seen)
    elif kind in {"bullets", "bullet", "list"}:
        items = [i for i in as_list(section.get("items") or section.get("body")) if plain_text(i)]
        if not items:
            return
        if heading:
            doc.add_heading(heading, level=3)
        add_bullets(doc, items)
    elif kind == "table":
        rows = []
        for row in as_list(section.get("rows")):
            if isinstance(row, dict):
                rows.append(
                    (
                        plain_text(row.get("key") or row.get("name") or row.get("项目")),
                        plain_text(row.get("value") or row.get("description") or row.get("说明")),
                    )
                )
            elif isinstance(row, (list, tuple)) and len(row) >= 2:
                rows.append((plain_text(row[0]), plain_text(row[1])))
        rows = [(k, v) for k, v in rows if plain_text(k) or plain_text(v)]
        if not rows:
            return
        if heading:
            doc.add_heading(heading, level=3)
        add_key_value_table(doc, rows)
    else:
        body = section.get("body")
        if body is None:
            body = section.get("items")
        paragraphs = [p for p in as_list(body) if plain_text(p)]
        if not paragraphs:
            return
        if heading:
            doc.add_heading(heading, level=3)
        for paragraph in paragraphs:
            add_paragraph(doc, paragraph)


def add_legacy_module_sections(
    doc: Document,
    module: dict[str, Any],
    base_dir: Path,
    state: DocState,
    seen: set[str],
) -> None:
    if [i for i in as_list(module.get("scenarios")) if isinstance(i, str) and i.strip()]:
        doc.add_heading("适用场景", level=3)
        add_bullets(doc, module.get("scenarios"))

    if clean_steps(module.get("steps")):
        doc.add_heading("操作步骤", level=3)
        add_steps(doc, clean_steps(module.get("steps")), base_dir, state, seen)

    if [i for i in as_list(module.get("key_outputs")) if isinstance(i, str) and i.strip()]:
        doc.add_heading("重点结果", level=3)
        add_bullets(doc, module.get("key_outputs"))

    if [i for i in as_list(module.get("notes")) if isinstance(i, str) and i.strip()]:
        doc.add_heading("注意事项", level=3)
        add_bullets(doc, module.get("notes"))


def add_module_full(doc: Document, module: dict[str, Any], base_dir: Path, state: DocState) -> None:
    name = plain_text(module.get("name"))
    if not name:
        return
    seen: set[str] = set()
    doc.add_heading(name, level=2)
    add_paragraph(doc, plain_text(module.get("purpose")))

    sections = [section for section in as_list(module.get("sections")) if isinstance(section, dict)]
    if sections:
        for section in sections:
            add_section(doc, section, base_dir, state, seen)
    else:
        add_legacy_module_sections(doc, module, base_dir, state, seen)

    add_screenshots(doc, module.get("screenshots"), base_dir, state, seen)


def add_modules(doc: Document, modules: Any, base_dir: Path, state: DocState) -> None:
    for module in as_list(modules):
        if isinstance(module, dict):
            add_module_full(doc, module, base_dir, state)


def add_module_index(doc: Document, module: dict[str, Any]) -> None:
    """Query-level index entry: name + purpose only."""
    name = plain_text(module.get("name"))
    if not name:
        return
    doc.add_heading(name, level=2)
    add_paragraph(doc, plain_text(module.get("purpose")))


ENTRY_FRAGMENT_RE = re.compile(r"「([^」]+)」")


def module_covered_by_tasks(module_name: str, tasks: list[dict[str, Any]]) -> bool:
    """True when a task already carries this module's usage details.

    Explicit mapping wins: if any task provides ``covers_modules``, only
    those exact names count (entry text is NOT consulted). Otherwise, fall
    back to the legacy heuristic that matches module names inside「」in
    task entries (exact fragment match, no substring false positives).
    """
    has_explicit = any("covers_modules" in task and isinstance(task.get("covers_modules"), list) for task in tasks)
    if has_explicit:
        for task in tasks:
            for name in as_list(task.get("covers_modules")):
                if plain_text(name) == module_name:
                    return True
        return False
    for task in tasks:
        entry = plain_text(task.get("entry"))
        if not entry:
            continue
        for fragment in ENTRY_FRAGMENT_RE.findall(entry):
            if fragment == module_name:
                return True
    return False


# ---------------------------------------------------------------------------
# V2 tasks
# ---------------------------------------------------------------------------

PRIORITY_LABELS = {"core": "核心", "supporting": "辅助", "reference": "参考"}


def add_task_prerequisites(doc: Document, prerequisites: Any) -> None:
    items = [plain_text(i) for i in as_list(prerequisites) if plain_text(i)]
    if not items:
        return
    doc.add_heading("开始前需要", level=3)
    add_bullets(doc, items)


def add_task_common_problems(doc: Document, problems: Any) -> None:
    items = [
        item
        for item in as_list(problems)
        if plain_text(item.get("question") if isinstance(item, dict) else item)
    ]
    if not items:
        return
    doc.add_heading("常见问题", level=3)
    for item in items:
        if isinstance(item, dict):
            question = plain_text(item.get("question"))
            answer = plain_text(item.get("answer"))
        else:
            question = plain_text(item)
            answer = ""
        if not question:
            continue
        q_para = doc.add_paragraph()
        q_para.paragraph_format.keep_with_next = True
        q_run = q_para.add_run(f"Q：{question}")
        q_run.bold = True
        set_run_font(q_run)
        if answer:
            a_para = doc.add_paragraph()
            a_para.paragraph_format.left_indent = Inches(0.2)
            a_run = a_para.add_run(answer)
            set_run_font(a_run)


def add_task(doc: Document, task: dict[str, Any], base_dir: Path, state: DocState) -> None:
    title = plain_text(task.get("title"))
    if not title:
        return
    seen: set[str] = set()
    priority = text(task.get("priority")).lower()
    tag = PRIORITY_LABELS.get(priority, "")

    heading = doc.add_heading(level=2)
    heading_run = heading.add_run(title)
    set_run_font(heading_run)
    if tag:
        tag_run = heading.add_run(f"（{tag}）")
        tag_run.font.size = Pt(10)
        tag_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        set_run_font(tag_run)

    entry = plain_text(task.get("entry"))
    purpose = plain_text(task.get("purpose"))
    steps = clean_steps(task.get("steps"))

    if entry:
        entry_para = doc.add_paragraph()
        entry_para.paragraph_format.keep_with_next = True
        entry_run = entry_para.add_run(f"入口：{entry}")
        entry_run.bold = True
        set_run_font(entry_run)
    if purpose:
        add_paragraph(doc, purpose)

    if priority == "reference":
        # Lookup-level information only: purpose, entry, short steps.
        if steps:
            doc.add_heading("操作步骤", level=3)
            add_steps(doc, steps, base_dir, state, seen, with_screenshots=False)
    elif priority == "supporting":
        add_task_prerequisites(doc, task.get("prerequisites"))
        if steps:
            doc.add_heading("操作步骤", level=3)
            add_steps(doc, steps, base_dir, state, seen)
        # Never silently drop explicitly provided result / common problems.
        result = task.get("result")
        result_items = [r for r in as_list(result) if plain_text(r)]
        if result_items:
            doc.add_heading("完成后你会看到", level=3)
            if isinstance(result, list):
                add_bullets(doc, result_items)
            else:
                add_paragraph(doc, result)
        add_task_common_problems(doc, task.get("common_problems"))
        add_screenshots(doc, task.get("screenshots"), base_dir, state, seen)
    else:
        # core; also the default for missing/unknown priority so no content
        # is silently dropped.
        add_task_prerequisites(doc, task.get("prerequisites"))
        if steps:
            doc.add_heading("操作步骤", level=3)
            add_steps(doc, steps, base_dir, state, seen)
        result = task.get("result")
        result_items = [r for r in as_list(result) if plain_text(r)]
        if result_items:
            doc.add_heading("完成后你会看到", level=3)
            if isinstance(result, list):
                add_bullets(doc, result_items)
            else:
                add_paragraph(doc, result)
        add_task_common_problems(doc, task.get("common_problems"))
        add_screenshots(doc, task.get("screenshots"), base_dir, state, seen)


def add_faq(doc: Document, faq: Any) -> None:
    items = as_list(faq)
    if not items:
        return
    doc.add_heading("常见问题", level=1)
    for item in items:
        if isinstance(item, dict):
            question = plain_text(item.get("question"))
            answer = plain_text(item.get("answer"))
        else:
            question = plain_text(item)
            answer = ""
        if question:
            doc.add_heading(question, level=3)
        if answer:
            add_paragraph(doc, answer)


# ---------------------------------------------------------------------------
# First-page 3-minute summary
# ---------------------------------------------------------------------------


def add_first_page(
    doc: Document,
    data: dict[str, Any],
    at_a_glance: dict[str, Any],
    overview: str,
    quick_start: list[str],
    tasks: list[dict[str, Any]],
    modules: list[dict[str, Any]],
    usage_rules: list[str],
    base_dir: Path,
    state: DocState,
) -> dict[str, Any]:
    what = plain_text(at_a_glance.get("what_it_does")) or overview
    top_tasks: list[str] = []
    seen_top: set[str] = set()
    for item in as_list(at_a_glance.get("top_tasks")):
        task_name = plain_text(item)
        if task_name and task_name not in seen_top:
            seen_top.add(task_name)
            top_tasks.append(task_name)
        if len(top_tasks) >= 5:
            break
    if not top_tasks:
        # Task-oriented fallback: task titles first, then module names.
        seen_titles: set[str] = set()
        for task in tasks:
            title = plain_text(task.get("title"))
            if title and title not in seen_titles:
                seen_titles.add(title)
                top_tasks.append(title)
            if len(top_tasks) >= 5:
                break
    if not top_tasks:
        seen_modules: set[str] = set()
        for module in modules:
            name = plain_text(module.get("name"))
            if name and name not in seen_modules:
                seen_modules.add(name)
                top_tasks.append(name)
            if len(top_tasks) >= 5:
                break
    before_start = [plain_text(b) for b in as_list(at_a_glance.get("before_you_start")) if plain_text(b)][:4]
    if not before_start and usage_rules:
        before_start = usage_rules[:4]
    access_url = plain_text(data.get("access_url"))

    page1_had_content = bool(what or top_tasks or quick_start or before_start or access_url)
    if not page1_had_content:
        return {"page1_had_content": False}

    doc.add_heading("3 分钟了解这个产品", level=1)
    if what:
        doc.add_heading("这个产品能帮你做什么", level=2)
        add_paragraph(doc, what)
    if top_tasks:
        doc.add_heading("最常用的几件事", level=2)
        add_bullets(doc, top_tasks)
    if quick_start:
        doc.add_heading("第一次建议这样用", level=2)
        add_steps(doc, quick_start, base_dir, state, set())
    if before_start:
        doc.add_heading("开始前要知道", level=2)
        add_bullets(doc, before_start)
    if access_url:
        doc.add_heading("系统入口", level=2)
        add_paragraph(doc, access_url)

    return {"page1_had_content": True, "before_start_used": before_start, "what_used": what}


def _steps_overlap(a: list[str], b: list[str]) -> bool:
    """True when two step lists clearly duplicate: equal, or one is fully
    contained as a prefix of the other."""
    if not a or not b:
        return False
    if a == b:
        return True
    shorter, longer = (a, b) if len(a) <= len(b) else (b, a)
    return longer[: len(shorter)] == shorter


# ---------------------------------------------------------------------------
# Main render
# ---------------------------------------------------------------------------


def render(data: dict[str, Any], output_path: Path, base_dir: Path) -> None:
    doc = Document()
    configure_document(doc)
    state = DocState()

    title = plain_text(data.get("title")) or "产品使用说明总结"
    title_para = doc.add_paragraph(style="Title")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    set_run_font(run)

    subtitle_parts = [
        plain_text(data.get("subtitle")),
        f"适用对象：{plain_text(data.get('audience'))}" if plain_text(data.get("audience")) else "",
        f"版本日期：{plain_text(data.get('version_date'))}" if plain_text(data.get("version_date")) else "",
    ]
    subtitle = "    ".join(part for part in subtitle_parts if part)
    if subtitle:
        para = doc.add_paragraph(style="Subtitle")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(subtitle)
        set_run_font(run)

    at_a_glance = data.get("at_a_glance") if isinstance(data.get("at_a_glance"), dict) else {}
    overview = plain_text(data.get("overview"))
    quick_start = [plain_text(q) for q in as_list(data.get("quick_start")) if plain_text(q)]
    recommended = [plain_text(r) for r in as_list(data.get("recommended_workflow")) if plain_text(r)]
    # If one clearly duplicates the other (equal or prefix), keep the longer
    # list as Quick Start and drop the duplicate Recommended Workflow.
    if recommended and _steps_overlap(recommended, quick_start):
        if len(recommended) > len(quick_start):
            quick_start = recommended
        recommended = []
    tasks = [t for t in as_list(data.get("tasks")) if isinstance(t, dict) and plain_text(t.get("title"))]
    modules = [m for m in as_list(data.get("modules")) if isinstance(m, dict) and plain_text(m.get("name"))]
    usage_rules = [plain_text(u) for u in as_list(data.get("usage_rules")) if plain_text(u)]
    faq_items = [
        item
        for item in as_list(data.get("faq"))
        if plain_text(item.get("question") if isinstance(item, dict) else item)
        or (isinstance(item, dict) and plain_text(item.get("answer")))
    ]

    page1 = add_first_page(
        doc, data, at_a_glance, overview, quick_start, tasks, modules, usage_rules, base_dir, state
    )

    metadata_rows = [
        ("使用对象", plain_text(data.get("audience"))),
        ("文档范围", plain_text(data.get("scope"))),
        ("额度/账号规则", plain_text(data.get("quota_rules") or data.get("account_rules"))),
        ("输出位置", plain_text(data.get("output_location"))),
    ]
    metadata_rows = [(k, v) for k, v in metadata_rows if plain_text(v)]
    # Page 1 already shows what_it_does (or its overview fallback); a
    # standalone 简短说明 only makes sense when overview adds different text.
    standalone_overview = bool(overview and overview != page1.get("what_used", ""))
    # Recommended Workflow was already deduplicated against Quick Start above.
    standalone_workflow = bool(recommended)
    # 使用规范 renders only the rules not already shown on page 1, so the
    # fallback (usage_rules -> 开始前要知道) never duplicates content.
    page1_rules_shown = set(page1.get("before_start_used") or [])
    standalone_rules_items = [u for u in usage_rules if u not in page1_rules_shown]
    standalone_rules = bool(standalone_rules_items)
    follows = bool(
        metadata_rows
        or standalone_overview
        or standalone_workflow
        or tasks
        or modules
        or standalone_rules
        or faq_items
    )
    if page1["page1_had_content"] and follows:
        doc.add_page_break()

    add_key_value_table(doc, metadata_rows)

    if standalone_overview:
        doc.add_heading("简短说明", level=1)
        add_paragraph(doc, overview)

    if standalone_workflow:
        doc.add_heading("推荐使用流程", level=1)
        add_steps(doc, recommended, base_dir, state, set())

    if tasks:
        doc.add_heading("常用任务", level=1)
        for task in tasks:
            add_task(doc, task, base_dir, state)

    if modules:
        if tasks:
            doc.add_heading("功能索引", level=1)
            for module in modules:
                name = plain_text(module.get("name"))
                if module_covered_by_tasks(name, tasks):
                    # Task sections already carry the details; index entry
                    # avoids duplicating content.
                    add_module_index(doc, module)
                else:
                    # Not covered by any task: render in full so no content
                    # is silently lost.
                    add_module_full(doc, module, base_dir, state)
        else:
            doc.add_heading("模块使用说明", level=1)
            add_modules(doc, modules, base_dir, state)

    if standalone_rules:
        doc.add_heading("使用规范", level=1)
        add_bullets(doc, standalone_rules_items)

    add_faq(doc, faq_items)

    add_header_footer(doc, title, plain_text(data.get("version_date")))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Render product usage summary JSON to DOCX.")
    parser.add_argument("--input", required=True, help="Path to summary JSON.")
    parser.add_argument("--output", required=True, help="Path to output DOCX.")
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    output_path = Path(args.output).resolve()
    data = json.loads(input_path.read_text(encoding="utf-8"))
    render(data, output_path, input_path.parent)


if __name__ == "__main__":
    main()
